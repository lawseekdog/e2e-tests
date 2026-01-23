"""DOCX download + content assertions."""

from __future__ import annotations

from io import BytesIO
from typing import Iterable


def extract_docx_text(docx_bytes: bytes) -> str:
    """Best-effort extraction of visible text from a docx file."""
    if not isinstance(docx_bytes, (bytes, bytearray)) or not docx_bytes:
        return ""
    from docx import Document  # python-docx

    doc = Document(BytesIO(docx_bytes))
    parts: list[str] = []

    for p in doc.paragraphs or []:
        if p is None:
            continue
        t = (p.text or "").strip()
        if t:
            parts.append(t)

    for tbl in doc.tables or []:
        for row in tbl.rows or []:
            for cell in row.cells or []:
                t = (cell.text or "").strip()
                if t:
                    parts.append(t)

    return "\n".join(parts)


def assert_docx_contains(text: str, *, must_include: Iterable[str]) -> None:
    missing: list[str] = []
    for needle in must_include:
        s = str(needle or "").strip()
        if not s:
            continue
        if s not in text:
            missing.append(s)
    if missing:
        sample = text[:2000]
        raise AssertionError(f"DOCX missing required fragments: {missing}. Extracted sample:\n{sample}")

